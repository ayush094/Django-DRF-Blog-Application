# from rest_framework import viewsets, generics, status, filters
# from rest_framework.views import APIView
# from rest_framework.response import Response
# from rest_framework.permissions import AllowAny, IsAuthenticated
# import logging
# from django.contrib.auth import authenticate
# from rest_framework.authtoken.models import Token
# from rest_framework.decorators import action
# from django.shortcuts import get_object_or_404
# from django_filters.rest_framework import DjangoFilterBackend
# from django.db.models import Q

# from .models import Blog, Comment, Like
# from .serializers import (
#     RegisterSerializer, UserSerializer,
#     BlogSerializer, BlogCreateSerializer, CommentSerializer
# )
# from .permissions import IsAuthorOrReadOnly
# from drf_yasg.utils import swagger_auto_schema
# from drf_yasg import openapi
# class RegisterAPI(generics.CreateAPIView):
#     permission_classes = [AllowAny]
#     serializer_class = RegisterSerializer

# class LoginAPI(APIView):
#     permission_classes = [AllowAny]
#     @swagger_auto_schema(
#         request_body=openapi.Schema(
#             type=openapi.TYPE_OBJECT,
#             required=['username', 'password'],
#             properties={
#                 'username': openapi.Schema(type=openapi.TYPE_STRING, example='alice'),
#                 'password': openapi.Schema(type=openapi.TYPE_STRING, example='StrongPass123'),
#             },
#         ),
#         responses={200: "Login successful"},
#     )
#     def post(self, request, *args, **kwargs):
#         username = request.data.get("username")
#         password = request.data.get("password")
        
#         if username is None or password is None:
#             return Response({"error": "Please provide both username and password."},
#                             status=status.HTTP_400_BAD_REQUEST)
        
#         user = authenticate(username=username, password=password)
        
#         if not user:
#             return Response({"error": "Invalid Credentials"},
#                             status=status.HTTP_401_UNAUTHORIZED)
#         token, created = Token.objects.get_or_create(user=user)
#         user_data = UserSerializer(user).data
#         return Response({"token": token.key , "user": user_data}, status=status.HTTP_200_OK)
    


# class LogoutAPI(APIView):
#     permission_classes = [IsAuthenticated]

#     def post(self, request, *args, **kwargs):
#         request.user.auth_token.delete()
#         return Response({'success': 'Logged out'}, status=status.HTTP_200_OK)

# class BlogViewSet(viewsets.ModelViewSet):
#     queryset = Blog.objects.all().order_by('-created_at')
#     serializer_class = BlogSerializer
#     permission_classes = [IsAuthorOrReadOnly]
#     filter_backends = [DjangoFilterBackend, filters.SearchFilter]
#     filterset_fields = ['author__username']
#     search_fields = ['title', 'content']

#     def get_queryset(self):
#         qs = Blog.objects.all().order_by('-created_at')
#         user = getattr(self.request, 'user', None)
#         if not user or not user.is_authenticated:
#             return qs.filter(is_published=True)
#         try:
#             if user.role == 'author':
#                 return qs.filter(Q(is_published=True) | Q(author=user))
#         except Exception:
#             pass
#         return qs.filter(is_published=True)

#     def get_serializer_class(self):
#         if self.action in ['create', 'update', 'partial_update']:
#             return BlogCreateSerializer
#         return BlogSerializer

#     def perform_create(self, serializer):
#         serializer.save(author=self.request.user)

#     def create(self, request, *args, **kwargs):
#         logger = logging.getLogger(__name__)
#         serializer = self.get_serializer(data=request.data)
#         if not serializer.is_valid():
#             logger.error("Blog create validation errors: %s", serializer.errors)
#             return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
#         self.perform_create(serializer)
#         headers = self.get_success_headers(serializer.data)
#         return Response(serializer.data, status=status.HTTP_201_CREATED, headers=headers)

#     @swagger_auto_schema(
#         method='post',
#         request_body=None,
#         responses={201: 'Liked'}
#     )
#     @action(detail=True, methods=['post'], permission_classes=[IsAuthenticated])
#     def like(self, request, pk=None):
#         blog = self.get_object()
#         user = request.user
#         if not blog.is_published and blog.author != user:
#             return Response({'detail': 'Cannot like an unpublished blog.'}, status=status.HTTP_403_FORBIDDEN)
#         like, created = Like.objects.get_or_create(blog=blog, user=user)
#         if not created:
#             return Response({'detail': 'Already liked.'}, status=status.HTTP_400_BAD_REQUEST)
#         serializer = self.get_serializer(blog, context={'request': request})
#         return Response(serializer.data, status=status.HTTP_201_CREATED)

#     @swagger_auto_schema(
#         method='post',
#         request_body=None,
#         responses={200: 'Unliked'}
#     )
#     @action(detail=True, methods=['post'], permission_classes=[IsAuthenticated])
#     def unlike(self, request, pk=None):
#         blog = self.get_object()
#         user = request.user
#         try:
#             like = Like.objects.get(blog=blog, user=user)
#             like.delete()
#             serializer = self.get_serializer(blog, context={'request': request})
#             return Response(serializer.data, status=status.HTTP_200_OK)
#         except Like.DoesNotExist:
#             return Response({'detail': 'Like not found.'}, status=status.HTTP_400_BAD_REQUEST)

#     @swagger_auto_schema(
#         method='post',
#         request_body=None,
#         responses={200: 'Published'}
#     )
#     @action(detail=True, methods=['post'], permission_classes=[IsAuthenticated, IsAuthorOrReadOnly])
#     def publish(self, request, pk=None):
#         blog = self.get_object()
#         if blog.author != request.user:
#             return Response({'detail': 'Only the author can publish this blog.'}, status=status.HTTP_403_FORBIDDEN)
#         if blog.is_published:
#             return Response({'detail': 'Blog already published.'}, status=status.HTTP_400_BAD_REQUEST)
#         blog.is_published = True
#         blog.save()
#         serializer = self.get_serializer(blog, context={'request': request})
#         return Response(serializer.data, status=status.HTTP_200_OK)

#     @swagger_auto_schema(
#         method='post',
#         request_body=None,
#         responses={200: 'Unpublished'}
#     )
#     @action(detail=True, methods=['post'], permission_classes=[IsAuthenticated, IsAuthorOrReadOnly])
#     def unpublish(self, request, pk=None):
#         blog = self.get_object()
#         if blog.author != request.user:
#             return Response({'detail': 'Only the author can unpublish this blog.'}, status=status.HTTP_403_FORBIDDEN)
#         if not blog.is_published:
#             return Response({'detail': 'Blog already unpublished.'}, status=status.HTTP_400_BAD_REQUEST)
#         blog.is_published = False
#         blog.save()
#         serializer = self.get_serializer(blog, context={'request': request})
#         return Response(serializer.data, status=status.HTTP_200_OK)


# class CommentCreateAPI(generics.CreateAPIView):
#     serializer_class = CommentSerializer
#     permission_classes = [IsAuthenticated]

#     def create(self, request, *args, **kwargs):
#         logger = logging.getLogger(__name__)
#         logger.info("CommentCreateAPI.create called with data: %s", request.data)
#         serializer = self.get_serializer(data=request.data)
#         if not serializer.is_valid():
#             logger.error("Comment create validation errors: %s", serializer.errors)
#             return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
#         self.perform_create(serializer)
#         headers = self.get_success_headers(serializer.data)
#         return Response(serializer.data, status=status.HTTP_201_CREATED, headers=headers)

#     def perform_create(self, serializer):
#         blog_id = self.kwargs.get('blog_id')
#         if blog_id:
#             blog = get_object_or_404(Blog, id=blog_id)
#             serializer.save(commenter=self.request.user, blog=blog)
#         else:
#             serializer.save(commenter=self.request.user)


    
# class BlogCommentsListAPI(generics.ListAPIView):
#     serializer_class = CommentSerializer
#     permission_classes = [AllowAny]

#     def get_queryset(self):
#         blog_id = self.kwargs.get('blog_id')
#         return Comment.objects.filter(blog__id=blog_id).order_by('-created_at')



# from django.shortcuts import get_object_or_404
# from rest_framework.decorators import api_view, permission_classes
# from rest_framework.permissions import AllowAny, IsAuthenticated
# from rest_framework.response import Response
# from rest_framework import status
# from rest_framework_simplejwt.tokens import RefreshToken
# from .models import Blog, Comment, Like
# from .serializers import (
#     BlogSerializer, BlogCreateSerializer,
#     CommentSerializer, UserSerializer, RegisterSerializer
# )
# from django.contrib.auth import authenticate
# from drf_yasg.utils import swagger_auto_schema
# from drf_yasg import openapi



# @swagger_auto_schema(
#     method='post',
#     request_body=openapi.Schema(
#         type=openapi.TYPE_OBJECT,
#         required=['username', 'password'],
#         properties={
#             'username': openapi.Schema(type=openapi.TYPE_STRING, example='ayush1234'),
#             'password': openapi.Schema(type=openapi.TYPE_STRING, example='ayush@1234'),
#         },
#     ),
#     responses={200: "Login successful"}
# )
# @api_view(['POST'])
# @permission_classes([AllowAny])
# def login(request):
#     username = request.data.get("username")
#     password = request.data.get("password")

#     user = authenticate(username=username, password=password)
#     if not user:
#         return Response({"error": "Invalid credentials"}, status=401)

#     refresh = RefreshToken.for_user(user)
#     return Response({
#         "User Is Authenticated": True,
#         "refresh": str(refresh),
#         "access": str(refresh.access_token),
#         "user": UserSerializer(user).data
#     })


# @swagger_auto_schema(
#     method='post',
#     request_body=RegisterSerializer,
#     responses={201: "User registered successfully"}
# )
# @api_view(['POST'])
# @permission_classes([AllowAny])
# def register(request):
#     serializer = RegisterSerializer(data=request.data)
#     if serializer.is_valid():
#         user = serializer.save()
#         return Response(UserSerializer(user).data, status=201)
#     return Response(serializer.errors, status=400)


# @swagger_auto_schema(
#     method='post',
#     security=[{"Bearer": []}]
# )
# @api_view(['POST'])
# @permission_classes([IsAuthenticated])
# def logout(request):
#     try:
#         refresh_token = request.data.get("refresh")
#         token = RefreshToken(refresh_token)
#         token.blacklist()
#         return Response({"success": "Logged out"}, status=200)
#     except:
#         return Response({"error": "Invalid token"}, status=400)


# @swagger_auto_schema(method="get", security=[{"Bearer": []}], responses={200: BlogSerializer(many=True)})
# @swagger_auto_schema(method="post", security=[{"Bearer": []}], request_body=BlogCreateSerializer)
# @api_view(['GET', 'POST'])
# @permission_classes([IsAuthenticated])
# def blog_list_create(request):
#     if request.method == 'GET':
#         blogs = Blog.objects.filter(is_published=True).order_by("-created_at")
#         return Response({
#             "User Is Authenticated": True,
#             "blogs":
#             BlogSerializer(blogs, many=True).data})

#     if request.method == 'POST':
#         serializer = BlogCreateSerializer(data=request.data)
#         if serializer.is_valid():
#             serializer.save(author=request.user)
#             return Response(serializer.data, status=201)
#         return Response(serializer.errors, status=400)



# @swagger_auto_schema(method='get', security=[{"Bearer": []}])
# @swagger_auto_schema(method='put', security=[{"Bearer": []}], request_body=BlogCreateSerializer)
# @swagger_auto_schema(method='patch', security=[{"Bearer": []}], request_body=BlogCreateSerializer)
# @swagger_auto_schema(method='delete', security=[{"Bearer": []}])
# @api_view(['GET', 'PUT', 'PATCH', 'DELETE'])
# @permission_classes([IsAuthenticated])
# def blog_detail(request, pk):
#     blog = get_object_or_404(Blog, pk=pk)

#     if request.method == 'GET':
#         if not blog.is_published and blog.author != request.user:
#             return Response({"detail": "Not published"}, status=403)
        
        
#         return Response({
#             "User Is Authenticated": True,
#             "blog":BlogSerializer(blog).data})

#     if blog.author != request.user:
#         return Response({"detail": "Not allowed"}, status=403)

#     if request.method in ['PUT', 'PATCH']:
#         serializer = BlogCreateSerializer(blog, data=request.data, partial=(request.method == 'PATCH'))
#         if serializer.is_valid():
#             serializer.save()
#             return Response(serializer.data)
#         return Response(serializer.errors, status=400)

#     blog.delete()
#     return Response({"detail": "Deleted"}, status=204)


# @swagger_auto_schema(method='post', security=[{"Bearer": []}])
# @api_view(['POST'])
# @permission_classes([IsAuthenticated])
# def blog_publish(request, pk):
#     blog = get_object_or_404(Blog, pk=pk)
#     if blog.author != request.user:
#         return Response({"detail": "Not allowed"}, status=403)
#     if blog.is_published:
#         return Response({"detail": "Already published"}, status=400)

#     blog.is_published = True
#     blog.save()
#     return Response(BlogSerializer(blog).data)


# @swagger_auto_schema(method='post', security=[{"Bearer": []}])
# @api_view(['POST'])
# @permission_classes([IsAuthenticated])
# def blog_unpublish(request, pk):
#     blog = get_object_or_404(Blog, pk=pk)
#     if blog.author != request.user:
#         return Response({"detail": "Not allowed"}, status=403)
#     if not blog.is_published:
#         return Response({"detail": "Already unpublished"}, status=400)

#     blog.is_published = False
#     blog.save()
#     return Response(BlogSerializer(blog).data)




# @swagger_auto_schema(method='post', security=[{"Bearer": []}])
# @api_view(['POST'])
# @permission_classes([IsAuthenticated])
# def like_blog(request, pk):
#     blog = get_object_or_404(Blog, pk=pk)
#     like, created = Like.objects.get_or_create(blog=blog, user=request.user)
#     if not created:
#         return Response({"detail": "Already liked"}, status=400)
#     return Response(BlogSerializer(blog).data, status=201)


# @swagger_auto_schema(method='post', security=[{"Bearer": []}])
# @api_view(['POST'])
# @permission_classes([IsAuthenticated])
# def unlike_blog(request, pk):
#     blog = get_object_or_404(Blog, pk=pk)
#     try:
#         Like.objects.get(blog=blog, user=request.user).delete()
#         return Response(BlogSerializer(blog).data)
#     except Like.DoesNotExist:
#         return Response({"detail": "Not liked"}, status=400)


# @swagger_auto_schema(method='post', request_body=CommentSerializer, security=[{"Bearer": []}])
# @api_view(['POST'])
# @permission_classes([IsAuthenticated])
# def create_comment(request, blog_id):
#     blog = get_object_or_404(Blog, id=blog_id)
#     serializer = CommentSerializer(data=request.data)
#     if serializer.is_valid():
#         serializer.save(commenter=request.user, blog=blog)
#         return Response(serializer.data, status=201)
#     return Response(serializer.errors, status=400)


# @swagger_auto_schema(method='get', security=[{"Bearer": []}])
# @api_view(['GET'])
# @permission_classes([AllowAny])
# def list_comments(request, blog_id):
#     comments = Comment.objects.filter(blog__id=blog_id)
#     return Response(CommentSerializer(comments, many=True).data)


# blog_app/views.py
# views.py (FULL)
from io import BytesIO
import os
import zipfile
import re
from dateutil import parser

from django.contrib.auth import authenticate
from django.contrib.auth.models import AnonymousUser
from django.shortcuts import get_object_or_404
from django.utils import timezone
from django.utils.text import slugify
from django.http import HttpResponse, FileResponse

from rest_framework import viewsets, permissions, filters, status
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.pagination import PageNumberPagination

from drf_yasg.utils import swagger_auto_schema
from drf_yasg import openapi

from rest_framework_simplejwt.tokens import RefreshToken

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.utils import ImageReader

from myproject.mqtt_client import publish_message

from .models import Blog, Comment, Like
from .serializers import (
    BlogSerializer,
    BlogCreateSerializer,
    CommentSerializer,
    UserSerializer,
    RegisterSerializer,
)
from .tasks import publish_scheduled_blog

from rest_framework.parsers import MultiPartParser, FormParser
# from myproject.mqtt_client import publish_message

from blog_app.mqtt_publisher import mqtt_publish
from django.db import connection


# -----------------------
# Pagination
# -----------------------
class DynamicPageNumberPagination(PageNumberPagination):
    page_size = 4
    page_size_query_param = "Number_of_Blog"
    max_page_size = 100


# -----------------------
# Auth ViewSet
# -----------------------
class AuthViewSet(viewsets.ViewSet):
    permission_classes = [permissions.AllowAny]

    @swagger_auto_schema(
        request_body=openapi.Schema(
            type=openapi.TYPE_OBJECT,
            required=["username", "password"],
            properties={
                "username": openapi.Schema(type=openapi.TYPE_STRING, example="ayush1234"),
                "password": openapi.Schema(type=openapi.TYPE_STRING, example="ayush@1234"),
            },
        ),
        responses={200: "Login successful", 401: "Invalid credentials"},
    )
    @action(detail=False, methods=["post"])
    def login(self, request):
        username = request.data.get("username")
        password = request.data.get("password")
        user = authenticate(username=username, password=password)
        if not user:
            return Response({"error": "Invalid credentials"}, status=status.HTTP_401_UNAUTHORIZED)
        refresh = RefreshToken.for_user(user)
        return Response(
            {
                "User Is Authenticated": True,
                "refresh": str(refresh),
                "access": str(refresh.access_token),
                "user": UserSerializer(user).data,
            }
        )

    @swagger_auto_schema(request_body=RegisterSerializer, responses={201: "User registered successfully", 400: "Bad request"})
    @action(detail=False, methods=["post"] ,   parser_classes=[MultiPartParser, FormParser])
    def register(self, request):
        """
        Register endpoint.
        Note: If you want authors to upload a logo during registration,
        ensure RegisterSerializer accepts 'logo' and saves it to the Author model.
        """
        serializer = RegisterSerializer(data=request.data)
        if serializer.is_valid():
            user = serializer.save()
            return Response(UserSerializer(user).data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(
        request_body=openapi.Schema(type=openapi.TYPE_OBJECT, properties={"refresh": openapi.Schema(type=openapi.TYPE_STRING)}),
        security=[{"Bearer": []}],
        responses={200: "Logged out", 400: "Invalid token"},
    )
    @action(detail=False, methods=["post"], permission_classes=[permissions.IsAuthenticated])
    def logout(self, request):
        """
        Logout by blacklisting refresh token.
        Expects JSON: {"refresh": "<refresh_token>"}
        """
        try:
            refresh_token = request.data.get("refresh")
            if not refresh_token:
                return Response({"error": "Refresh token required"}, status=status.HTTP_400_BAD_REQUEST)
            token = RefreshToken(refresh_token)
            token.blacklist()
            return Response({"success": "Logged out"}, status=status.HTTP_200_OK)
        except Exception:
            return Response({"error": "Invalid token"}, status=status.HTTP_400_BAD_REQUEST)


# -----------------------
# Blog ViewSet (FULL)
# -----------------------
class BlogViewSet(viewsets.ModelViewSet):
    
    filter_backends = [filters.SearchFilter]
    search_fields = ["title", "content"]
    pagination_class = DynamicPageNumberPagination
    permission_classes = [permissions.IsAuthenticated]

    def get_serializer_class(self):
        if self.action in ["create", "update", "partial_update"]:
            return BlogCreateSerializer
        return BlogSerializer


    def get_queryset(self):
        user = self.request.user

        queryset = (
            Blog.objects
            .select_related("author")                    
            .prefetch_related("comments", "likes")
            .order_by("-created_at")
        )

        # Normal users & anonymous → only published blogs
        if isinstance(user, AnonymousUser) or getattr(user, "role", None) == "user":
            return queryset.filter(is_published=True)

        print("DB queries so far:", len(connection.queries))
        
        
        # Author / admin → all blogs
        return queryset



    # -----------------------
    # CREATE (with scheduling)
    # -----------------------
    def create(self, request, *args, **kwargs):
        import json
        user = request.user

        if getattr(user, "role", None) != "author" and not user.is_staff:
            return Response({"detail": "Not allowed"}, status=status.HTTP_403_FORBIDDEN)

        data = request.data.copy()
        scheduled_time = data.get("scheduled_publish_at")
        data["is_published"] = False

        # Handle scheduling
        if scheduled_time:
            try:
                utc_time = parser.isoparse(scheduled_time).astimezone(timezone.utc)
                data["scheduled_publish_at"] = utc_time

                if utc_time <= timezone.now():
                    data["is_published"] = True
                    data["published_at"] = timezone.now()

            except Exception:
                return Response({"error": "Invalid datetime format"}, status=status.HTTP_400_BAD_REQUEST)
        else:
            data["scheduled_publish_at"] = None

        serializer = BlogCreateSerializer(data=data, context={'request': request})    
            
        if serializer.is_valid():
            blog = serializer.save(author=user)

            # -----------------------------------
            # CASE 1 → Scheduled blog (Future ETA)
            # -----------------------------------
            if blog.scheduled_publish_at and blog.scheduled_publish_at > timezone.now():
                publish_scheduled_blog.apply_async(args=[blog.id], eta=blog.scheduled_publish_at)

                # MQTT notify blog scheduled
                publish_message(
                    "blog/scheduled",
                    json.dumps({
                        "blog_id": blog.id,
                        "title": blog.title,
                        "scheduled_publish_at": str(blog.scheduled_publish_at),
                        "status": "scheduled"
                    })
                )

            else:
                # -----------------------------------
                # CASE 2 → Immediate publish
                # -----------------------------------
                publish_message(
                    "blog/new",
                    json.dumps({
                        "blog_id": blog.id,
                        "title": blog.title,
                        "status": "created",
                        "published_at": str(blog.published_at)
                    })
                )

            return Response(BlogSerializer(blog, context={"request": request}).data, status=status.HTTP_201_CREATED)

        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    @swagger_auto_schema(
        operation_summary="Get a specific blog",
        manual_parameters=[
            openapi.Parameter(
                'blog_password', 
                openapi.IN_QUERY, 
                description="Enter the secret code if the blog is password protected", 
                type=openapi.TYPE_STRING
            )
        ],
        responses={
            200: BlogSerializer(),
            403: "This blog is locked. Please provide the correct Blog Password.",
            404: "Not Found"
        }
    )

    def retrieve(self, request, *args, **kwargs):

        instance = self.get_object()

        # Logic: If the blog is password protected
        if instance.is_password_protected:
            # Check if the user sent the "blog_password" in the URL or headers
            provided_password = request.query_params.get('blog_password')

            # If the password doesn't match the one stored in the Blog model
            if provided_password != instance.password:
                return Response(
                    {"detail": "This blog is locked. Please provide the correct Blog Password."}, 
                    status=status.HTTP_403_FORBIDDEN
                )

        # If it matches or isn't protected, show the blog content
        serializer = self.get_serializer(instance)
        return Response(serializer.data)
    

    # -----------------------
    # UPDATE / DELETE (permission enforced)
    # -----------------------
    def update(self, request, *args, **kwargs):
        user = request.user
        if getattr(user, "role", None) != "author" and not user.is_staff:
            return Response({"detail": "Not allowed"}, status=status.HTTP_403_FORBIDDEN)
        return super().update(request, *args, **kwargs)

    def destroy(self, request, *args, **kwargs):
        user = request.user
        if getattr(user, "role", None) != "author" and not user.is_staff:
            return Response({"detail": "Not allowed"}, status=status.HTTP_403_FORBIDDEN)
        return super().destroy(request, *args, **kwargs)

    # -----------------------
    # PUBLISH / UNPUBLISH
    # -----------------------
    @action(detail=True, methods=["post"])
    def publish(self, request, pk=None):
        blog = self.get_object()
        user = request.user
        if blog.author != user and not user.is_staff:
            return Response({"detail": "Not allowed"}, status=status.HTTP_403_FORBIDDEN)
        blog.is_published = True
        blog.published_at = timezone.now()
        blog.save()
        return Response(BlogSerializer(blog, context={"request": request}).data)

    @action(detail=True, methods=["post"])
    def unpublish(self, request, pk=None):
        blog = self.get_object()
        user = request.user
        if blog.author != user and not user.is_staff:
            return Response({"detail": "Not allowed"}, status=status.HTTP_403_FORBIDDEN)
        blog.is_published = False
        blog.save()
        return Response(BlogSerializer(blog, context={"request": request}).data)

    # -----------------------
    # LIKE / UNLIKE
    # -----------------------
    @action(detail=True, methods=["post"])
    def like(self, request, pk=None):
        blog = self.get_object()
        like, created = Like.objects.get_or_create(blog=blog, user=request.user)
        if not created:
            return Response({"detail": "Already liked"}, status=status.HTTP_400_BAD_REQUEST)
        return Response({"detail": "Liked"}, status=status.HTTP_201_CREATED)

    @action(detail=True, methods=["post"])
    def unlike(self, request, pk=None):
        blog = self.get_object()
        try:
            Like.objects.get(blog=blog, user=request.user).delete()
            return Response({"detail": "Unliked"}, status=status.HTTP_200_OK)
        except Like.DoesNotExist:
            return Response({"detail": "Not liked"}, status=status.HTTP_400_BAD_REQUEST)

    # -----------------------
    # SCHEDULE PUBLISH (action)
    # -----------------------
    @swagger_auto_schema(
        request_body=openapi.Schema(
            type=openapi.TYPE_OBJECT,
            properties={'scheduled_publish_at': openapi.Schema(type=openapi.TYPE_STRING, example="2025-12-02T10:45:00.000Z")}
        ),
        security=[{"Bearer": []}],
    )
    @action(detail=True, methods=["post"], url_path="schedule-publish")
    def schedule_publish(self, request, pk=None):
        blog = self.get_object()
        user = request.user
        if blog.author != user and not user.is_staff:
            return Response({"error": "Not allowed"}, status=status.HTTP_403_FORBIDDEN)

        dt_str = request.data.get("scheduled_publish_at")
        if not dt_str:
            return Response({"error": "scheduled_publish_at required"}, status=status.HTTP_400_BAD_REQUEST)

        try:
            utc_time = parser.isoparse(dt_str).astimezone(timezone.utc)
        except Exception:
            return Response({"error": "Invalid datetime format"}, status=status.HTTP_400_BAD_REQUEST)

        if utc_time <= timezone.now():
            blog.is_published = True
            blog.published_at = timezone.now()
        else:
            blog.is_published = False
            blog.published_at = None

        blog.scheduled_publish_at = utc_time
        blog.save()

        if utc_time > timezone.now():
            publish_scheduled_blog.apply_async(args=[blog.id], eta=utc_time)

        return Response({
            "message": "Scheduled",
            "blog_id": blog.id,
            "scheduled_publish_at_UTC": utc_time,
        }, status=status.HTTP_200_OK)

    # ----------------------------------------------------------
    # HELPER FUNCTIONS (auto-step + cleanup + line wrapping)
    # ----------------------------------------------------------
    def _normalize(self, text):
        """
        Normalize whitespace and control characters:
         - convert CRLF -> LF
         - replace tabs with single spaces
         - collapse multiple spaces
         - collapse 3+ newlines to 2 newlines (keep paragraphs)
         - trim edges
        """
        if not text:
            return ""
        text = text.replace("\r\n", "\n").replace("\r", "\n").replace("\t", " ")
        text = re.sub(r" +", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _auto_stepify(self, text):
        """
        Convert raw DB content into a list of steps for exports.
        Priority:
          1) multiple paragraphs (blank-line separated) -> each paragraph is a step
          2) else split by sentence boundaries (., !, ?)
          3) else whole text as single step
        Does NOT modify DB — just returns steps for the PDF/DOCX builder.
        """
        txt = self._normalize(text)
        if not txt:
            return []

        # 1) Paragraphs as steps
        parts = [p.strip() for p in re.split(r"\n\s*\n", txt) if p.strip()]
        if len(parts) > 1:
            return parts

        # 2) Sentence split
        sentences = [s.strip() for s in re.split(r"(?<=[.!?])\s+", txt) if s.strip()]
        steps = []
        for s in sentences:
            if not s.endswith((".", "?", "!")):
                s = s + "."
            steps.append(s)
        if steps:
            return steps

        # 3) Fallback
        return [txt]

    def _wrap_text(self, text, max_chars=80):
        """
        Naive char-based wrapping that slices text into chunks of max_chars.
        This is intentionally simple; ReportLab will place each chunk as a separate line.
        """
        if text is None:
            return [""]
        text = text.strip()
        if not text:
            return [""]
        return [text[i:i+max_chars] for i in range(0, len(text), max_chars)]

    # ----------------------------------------------------------
    # SINGLE BLOG DOWNLOAD - DOCX (with author logo + auto-steps)
    # ----------------------------------------------------------
    @swagger_auto_schema(
        operation_summary="Download single published blog as DOCX",
        responses={200: "DOCX File", 403: "Blog not published", 404: "Blog not found"},
    )
    @action(detail=True, methods=["get"], url_path="download/docx", permission_classes=[permissions.AllowAny])
    def download_docx(self, request, pk=None):
        blog = self.get_object()
        if not blog.is_published:
            return Response({"error": "Blog is not published yet"}, status=status.HTTP_403_FORBIDDEN)

        document = Document()

        # Insert author logo at top if available (local filesystem)
        try:
            logo_field = getattr(blog.author, "logo", None)
            if logo_field and getattr(logo_field, "path", None) and os.path.exists(logo_field.path):
                # width set to ~1.2 inch
                document.add_picture(logo_field.path, width=Inches(1.2))
                document.add_paragraph()  # space after logo
        except Exception:
            # ignore logo problems
            pass

        # Title
        title = document.add_heading(blog.title or "Untitled", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if title.runs:
            r = title.runs[0]
            r.font.size = Pt(22)
            r.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

        # Metadata
        meta = document.add_paragraph()
        meta_text = f"Author: {getattr(blog.author, 'username', 'Unknown')}  |  Published: {blog.published_at or 'N/A'}"
        meta_run = meta.add_run(meta_text)
        meta_run.font.size = Pt(12)
        meta_run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

        document.add_paragraph()  # spacing

        # Steps from content (auto-stepify)
        steps = self._auto_stepify(blog.content or "")

        # Render steps
        for idx, block in enumerate(steps, start=1):
            if not block:
                continue
            # Step header
            p_h = document.add_paragraph()
            r_h = p_h.add_run(f"STEP {idx}")
            r_h.bold = True
            r_h.font.size = Pt(16)
            r_h.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

            # Step body (preserve bullets if any)
            for paragraph in (block or "").split("\n"):
                p_body = document.add_paragraph()
                r_body = p_body.add_run(paragraph.strip())
                r_body.font.size = Pt(12)

        # Return as file
        stream = BytesIO()
        document.save(stream)
        stream.seek(0)
        filename = f"{slugify(blog.title) or 'blog'}-{blog.id}.docx"
        return FileResponse(stream, as_attachment=True, filename=filename)

    # ----------------------------------------------------------
    # SINGLE BLOG DOWNLOAD - PDF (with author logo + auto-steps)
    # ----------------------------------------------------------
    @swagger_auto_schema(
        operation_summary="Download single published blog as PDF",
        responses={200: "PDF File", 403: "Blog not published", 404: "Blog not found"},
    )
    @action(detail=True, methods=["get"], url_path="download/pdf", permission_classes=[permissions.AllowAny])
    def download_pdf(self, request, pk=None):
        blog = self.get_object()
        if not blog.is_published:
            return Response({"error": "Blog is not published yet"}, status=status.HTTP_403_FORBIDDEN)

        buffer = BytesIO()
        p = canvas.Canvas(buffer, pagesize=letter)
        width, height = letter
        left_margin = 50
        right_margin = 50
        max_width = width - left_margin - right_margin
        y = height - 50
        line_height = 14
        wrap_chars = 70  # characters per line inside box

        # ------------ AUTHOR LOGO ------------
        try:
            logo_field = getattr(blog.author, "logo", None)
            if logo_field and getattr(logo_field, "path", None) and os.path.exists(logo_field.path):
                reader = ImageReader(logo_field.path)
                logo_w = 60
                logo_h = 60
                p.drawImage(reader, left_margin, y - logo_h, width=logo_w, height=logo_h, mask='auto')
                y -= (logo_h + 10)
        except Exception:
            # silent fail - continue without logo
            pass

        # ------------ TITLE ------------
        p.setFont("Helvetica-Bold", 20)
        p.setFillColor(colors.HexColor("#1A73E8"))
        p.drawCentredString(width / 2, y, blog.title or "Untitled")
        y -= 40

        # ------------ METADATA BOX ------------
        p.setFont("Helvetica", 11)
        meta_text = f"Author: {getattr(blog.author, 'username', 'Unknown')}  |  Published: {blog.published_at or 'N/A'}"
        p.setStrokeColor(colors.HexColor("#1A73E8"))
        # draw a simple box
        p.rect(left_margin, y - 18, max_width, 30, stroke=1, fill=0)
        p.setFillColor(colors.black)
        p.drawString(left_margin + 10, y, meta_text)
        y -= 50

        # ------------ CONTENT (auto-step) ------------
        content = (self._normalize(blog.content or "")).strip()
        steps = self._auto_stepify(content)

        # ---------- CASE: RENDER STEPS ----------
        for index, block in enumerate(steps, start=1):
            if not block:
                continue

            # page break safety for heading
            if y < 100:
                p.showPage()
                y = height - 50

            # Step header
            p.setFont("Helvetica-Bold", 14)
            p.setFillColor(colors.HexColor("#0D47A1"))
            p.drawString(left_margin, y, f"STEP {index}")
            y -= 20

            # Prepare wrapped lines
            raw_lines = block.split("\n")
            wrapped_lines = []
            for ln in raw_lines:
                ln = ln.strip()
                if not ln:
                    wrapped_lines.append("")
                else:
                    wrapped_lines += self._wrap_text(ln, wrap_chars)

            total_lines = len(wrapped_lines)
            rect_height = total_lines * line_height + 20

            # Page break if not enough space
            if y - rect_height < 70:
                p.showPage()
                y = height - 50

            # Draw step box
            p.setStrokeColor(colors.HexColor("#90CAF9"))
            p.setFillColor(colors.HexColor("#E3F2FD"))
            p.rect(left_margin, y - rect_height, max_width, rect_height, fill=1, stroke=1)

            # Draw text inside box
            p.setFillColor(colors.black)
            p.setFont("Helvetica", 11)
            text_y = y - 15
            for ln in wrapped_lines:
                if text_y < 70:
                    p.showPage()
                    text_y = height - 50
                p.drawString(left_margin + 10, text_y, ln)
                text_y -= line_height

            y = text_y - 20  # gap after box

        # End of PDF
        p.showPage()
        p.save()
        buffer.seek(0)

        filename = f"{slugify(blog.title) or 'blog'}-{blog.id}.pdf"
        return FileResponse(buffer, as_attachment=True, filename=filename, content_type="application/pdf")

    # ----------------------------------------------------------
    # BULK DOWNLOAD - DOCX (all published)
    # ----------------------------------------------------------
    @swagger_auto_schema(
        manual_parameters=[], 
        operation_summary="Download all published blogs as single DOCX",
        responses={200: "DOCX File"},
    )
    @action(detail=False, methods=["get"], url_path="bulk/download/docx", permission_classes=[permissions.AllowAny])
    def bulk_download_docx(self, request):
        document = Document()
        header = document.add_heading("Published Blogs Report", level=1)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if header.runs:
            header.runs[0].font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
            header.runs[0].font.size = Pt(18)

        blogs = Blog.objects.filter(is_published=True).order_by("-published_at")
        if not blogs.exists():
            document.add_paragraph("No published blogs found.")

        for blog in blogs:
            # Try add logo
            try:
                logo_field = getattr(blog.author, "logo", None)
                if logo_field and getattr(logo_field, "path", None) and os.path.exists(logo_field.path):
                    document.add_picture(logo_field.path, width=Inches(1.0))
            except Exception:
                pass

            title = document.add_heading(blog.title or "Untitled", level=2)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if title.runs:
                title.runs[0].font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

            meta = document.add_paragraph()
            meta_run = meta.add_run(f"Author: {getattr(blog.author, 'username', 'Unknown')}  |  Published: {blog.published_at or 'N/A'}")
            meta_run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

            document.add_paragraph()

            steps = self._auto_stepify(blog.content or "")
            if not steps:
                # fallback: print whole content
                p = document.add_paragraph()
                p.add_run((blog.content or "")[:1000]).font.size = Pt(12)
            else:
                for idx, block in enumerate(steps, start=1):
                    if not block:
                        continue
                    sh = document.add_paragraph()
                    r = sh.add_run(f"STEP {idx}")
                    r.bold = True
                    r.font.size = Pt(15)
                    r.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
                    for paragraph in block.split("\n"):
                        p = document.add_paragraph()
                        p.add_run(paragraph).font.size = Pt(12)

            document.add_page_break()

        out = BytesIO()
        document.save(out)
        out.seek(0)
        return FileResponse(out, as_attachment=True, filename="published_blogs.docx")

    # ----------------------------------------------------------
    # BULK DOWNLOAD - PDF (all published)
    # ----------------------------------------------------------
    @swagger_auto_schema(
        manual_parameters=[], 
        operation_summary="Download all published blogs as single PDF",
        responses={200: "PDF File"},
    )
    @action(detail=False, methods=["get"], url_path="bulk/download/pdf", permission_classes=[permissions.AllowAny])
    def bulk_download_pdf(self, request):
        buf = BytesIO()
        p = canvas.Canvas(buf, pagesize=letter)
        width, height = letter
        left_margin = 50
        right_margin = 50
        max_width = width - left_margin - right_margin
        y = height - 50
        line_height = 14
        wrap_chars = 70

        blogs = Blog.objects.filter(is_published=True).order_by("-published_at")
        if not blogs.exists():
            p.setFont("Helvetica", 12)
            p.drawString(left_margin, y, "No published blogs found.")
            p.showPage()
        else:
            for blog in blogs:
                # Try draw logo
                try:
                    logo_field = getattr(blog.author, "logo", None)
                    if logo_field and getattr(logo_field, "path", None) and os.path.exists(logo_field.path):
                        reader = ImageReader(logo_field.path)
                        logo_w = 48
                        logo_h = 48
                        p.drawImage(reader, left_margin, y - logo_h, width=logo_w, height=logo_h, mask='auto')
                        y -= (logo_h + 6)
                except Exception:
                    pass

                # Title
                p.setFont("Helvetica-Bold", 16)
                p.setFillColor(colors.HexColor("#1A73E8"))
                p.drawCentredString(width / 2, y, blog.title or "Untitled")
                y -= 30

                meta_text = f"Author: {getattr(blog.author, 'username', 'Unknown')}  |  Published: {blog.published_at or 'N/A'}"
                p.setFont("Helvetica", 10)
                p.rect(left_margin, y - 14, max_width, 24, stroke=1)
                p.setFillColor(colors.black)
                p.drawString(left_margin + 8, y, meta_text)
                y -= 34

                steps = self._auto_stepify(blog.content or "")

                if not steps:
                    # print fallback
                    for paragraph in (blog.content or "").split("\n"):
                        if not paragraph.strip():
                            y -= line_height
                            if y < 80:
                                p.showPage()
                                y = height - 50
                            continue
                        for chunk in self._wrap_text(paragraph, wrap_chars):
                            if y < 80:
                                p.showPage()
                                y = height - 50
                            p.setFont("Helvetica", 11)
                            p.drawString(left_margin, y, chunk)
                            y -= line_height
                else:
                    for idx, block in enumerate(steps, start=1):
                        if not block:
                            continue
                        if y < 110:
                            p.showPage()
                            y = height - 50

                        p.setFont("Helvetica-Bold", 13)
                        p.setFillColor(colors.HexColor("#0D47A1"))
                        p.drawString(left_margin, y, f"STEP {idx}")
                        y -= 18

                        raw_lines = block.split("\n")
                        wrapped_lines = []
                        for ln in raw_lines:
                            ln = ln.strip()
                            if not ln:
                                wrapped_lines.append("")
                            else:
                                wrapped_lines += self._wrap_text(ln, wrap_chars)

                        total_lines = len(wrapped_lines)
                        rect_h = total_lines * line_height + 12
                        if y - rect_h < 80:
                            p.showPage()
                            y = height - 50

                        p.setStrokeColor(colors.HexColor("#90CAF9"))
                        p.setFillColor(colors.HexColor("#E3F2FD"))
                        p.rect(left_margin, y - rect_h, max_width, rect_h + 6, fill=1, stroke=1)

                        p.setFillColor(colors.black)
                        p.setFont("Helvetica", 11)

                        temp_y = y - 8
                        for ln in wrapped_lines:
                            if not ln.strip():
                                temp_y -= line_height
                                if temp_y < 80:
                                    p.showPage()
                                    temp_y = height - 50
                                continue
                            if temp_y < 80:
                                p.showPage()
                                temp_y = height - 50
                            p.drawString(left_margin + 8, temp_y, ln)
                            temp_y -= line_height

                        y = temp_y - 16
                        if y < 80:
                            p.showPage()
                            y = height - 50

                p.showPage()
                y = height - 50

        p.save()
        buf.seek(0)
        return FileResponse(buf, as_attachment=True, filename="published_blogs.pdf", content_type="application/pdf")

    # ----------------------------------------------------------
    # BULK DOWNLOAD - ZIP (each published blog -> separate PDF inside zip)
    # ----------------------------------------------------------
    @swagger_auto_schema(
        manual_parameters=[],
        operation_summary="Download all published blogs as ZIP of PDFs",
        responses={200: "ZIP File"},
    )
    @action(detail=False, methods=["get"], url_path="bulk/download/zip", permission_classes=[permissions.AllowAny])
    def bulk_download_zip(self, request):
        blogs = Blog.objects.filter(is_published=True).order_by("-published_at")
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
            for blog in blogs:
                pdf_buf = BytesIO()
                p = canvas.Canvas(pdf_buf, pagesize=letter)
                width, height = letter
                left_margin = 50
                right_margin = 50
                max_width = width - left_margin - right_margin
                y = height - 50
                line_height = 14
                wrap_chars = 70

                # Title & metadata
                try:
                    logo_field = getattr(blog.author, "logo", None)
                    if logo_field and getattr(logo_field, "path", None) and os.path.exists(logo_field.path):
                        reader = ImageReader(logo_field.path)
                        p.drawImage(reader, left_margin, y-48, width=48, height=48, mask='auto')
                        y -= 56
                except Exception:
                    pass

                p.setFont("Helvetica-Bold", 14)
                p.drawString(left_margin, y, (blog.title or "Untitled")[:120])
                y -= 20

                p.setFont("Helvetica", 10)
                p.drawString(left_margin, y, f"Author: {getattr(blog.author, 'username', 'Unknown')}")
                y -= 12
                p.drawString(left_margin, y, f"Published At: {blog.published_at} | Scheduled: {blog.scheduled_publish_at}")
                y -= 12
                p.drawString(left_margin, y, f"Created: {blog.created_at} | Updated: {blog.updated_at}")
                y -= 14
                p.drawString(left_margin, y, f"Likes: {blog.likes.count()}  Comments: {blog.comments.count()}")
                y -= 16

                # Content
                steps = self._auto_stepify(blog.content or "")
                if not steps:
                    for paragraph in (blog.content or "").split("\n"):
                        paragraph = paragraph.strip()
                        if not paragraph:
                            y -= line_height
                            if y < 80:
                                p.showPage()
                                y = height - 50
                            continue
                        for chunk in self._wrap_text(paragraph, wrap_chars):
                            if y < 80:
                                p.showPage()
                                y = height - 50
                            p.drawString(left_margin, y, chunk)
                            y -= line_height
                else:
                    for idx, block in enumerate(steps, start=1):
                        if not block:
                            continue
                        if y < 110:
                            p.showPage()
                            y = height - 50

                        p.setFont("Helvetica-Bold", 13)
                        p.setFillColor(colors.HexColor("#0D47A1"))
                        p.drawString(left_margin, y, f"STEP {idx}")
                        y -= 18

                        raw_lines = block.split("\n")
                        wrapped_lines = []
                        for ln in raw_lines:
                            ln = ln.strip()
                            if not ln:
                                wrapped_lines.append("")
                            else:
                                wrapped_lines += self._wrap_text(ln, wrap_chars)

                        total_lines = len(wrapped_lines)
                        rect_h = total_lines * line_height + 12
                        if y - rect_h < 80:
                            p.showPage()
                            y = height - 50

                        p.setStrokeColor(colors.HexColor("#90CAF9"))
                        p.setFillColor(colors.HexColor("#E3F2FD"))
                        p.rect(left_margin, y - rect_h, max_width, rect_h + 6, fill=1, stroke=1)

                        p.setFillColor(colors.black)
                        p.setFont("Helvetica", 11)

                        temp_y = y - 8
                        for ln in wrapped_lines:
                            if not ln.strip():
                                temp_y -= line_height
                                if temp_y < 80:
                                    p.showPage()
                                    temp_y = height - 50
                                continue
                            if temp_y < 80:
                                p.showPage()
                                temp_y = height - 50
                            p.drawString(left_margin + 8, temp_y, ln)
                            temp_y -= line_height

                        y = temp_y - 16
                        if y < 80:
                            p.showPage()
                            y = height - 50

                p.showPage()
                p.save()
                pdf_buf.seek(0)
                filename = f"{slugify(blog.title) or 'blog'}-{blog.id}.pdf"
                zip_file.writestr(filename, pdf_buf.read())

        zip_buffer.seek(0)
        return FileResponse(zip_buffer, as_attachment=True, filename="published_blogs.zip")


# -----------------------
# Comment ViewSet
# -----------------------
class CommentViewSet(viewsets.ViewSet):
    permission_classes = [permissions.IsAuthenticated]

    @swagger_auto_schema(
        method='post',
        request_body=openapi.Schema(
            type=openapi.TYPE_OBJECT,
            required=['content'],
            properties={'content': openapi.Schema(type=openapi.TYPE_STRING, example="This is my comment")}
        ),
        responses={201: "Comment created successfully", 400: "Bad request"}
    )
    @action(detail=False, methods=["post"], url_path=r"create/(?P<blog_id>[^/.]+)")
    def create_comment(self, request, blog_id=None):
        blog = get_object_or_404(Blog, id=blog_id)

        content = request.data.get("content")
        if not content or content.strip() == "":
            return Response({"error": "Comment content is required"}, status=status.HTTP_400_BAD_REQUEST)

        serializer = CommentSerializer(data={
            "content": content,
            "blog": blog.id,
            "commenter": request.user.id
        })

        if serializer.is_valid():
            serializer.save(blog=blog, commenter=request.user)
            return Response({"message": "Comment created successfully", "data": serializer.data}, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    @action(detail=False, methods=["get"], url_path=r"list/(?P<blog_id>[^/.]+)")
    def list_comments(self, request, blog_id=None):
        comments = Comment.objects.filter(blog__id=blog_id).order_by("-created_at")
        paginator = DynamicPageNumberPagination()
        page = paginator.paginate_queryset(comments, request)

        if page is not None:
            return paginator.get_paginated_response(CommentSerializer(page, many=True).data)

        return Response(CommentSerializer(comments, many=True).data)